import pandas as pd
import streamlit as st
import io

from modules.lead_numeric.db import get_conn
from modules.lead_numeric.ddl import init_db

st.set_page_config(page_title="05 - Materialita", layout="wide")
st.title("05 - Materialita")
st.markdown(
    """
    <style>
    .block-container {padding-top: 1.8rem; padding-bottom: 0.6rem;}
    [data-testid="stVerticalBlock"] {gap: 0.35rem;}
    div[data-testid="stMetric"] {padding-top: 0.2rem; padding-bottom: 0.2rem;}
    [data-testid="stDataEditor"] [role="columnheader"][aria-colindex="3"],
    [data-testid="stDataEditor"] [role="columnheader"][aria-colindex="4"],
    [data-testid="stDataEditor"] [role="columnheader"][aria-colindex="5"],
    [data-testid="stDataEditor"] [role="gridcell"][aria-colindex="3"],
    [data-testid="stDataEditor"] [role="gridcell"][aria-colindex="4"],
    [data-testid="stDataEditor"] [role="gridcell"][aria-colindex="5"] {
        text-align: center !important;
        justify-content: center !important;
    }
    [data-testid="stDataEditor"] [role="gridcell"][aria-colindex="5"] input,
    [data-testid="stDataEditor"] [role="gridcell"][aria-colindex="3"] input,
    [data-testid="stDataEditor"] [role="gridcell"][aria-colindex="4"] input {
        text-align: center !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

init_db()

DEFAULT_ROWS = pd.DataFrame(
    [
        {"Voce": "A1) Ricavi delle vendite e delle prestazioni", "Valore base": 0, "% min": 1, "% max": 3, "% selezionata": 2, "Importo calcolato": 0, "Selezione": False},
        {"Voce": "Totale attivo", "Valore base": 0, "% min": 1, "% max": 3, "% selezionata": 2, "Importo calcolato": 0, "Selezione": False},
        {"Voce": "Patrimonio netto", "Valore base": 0, "% min": 3, "% max": 5, "% selezionata": 4, "Importo calcolato": 0, "Selezione": False},
        {"Voce": "Reddito ante imposte", "Valore base": 0, "% min": 3, "% max": 7, "% selezionata": 5, "Importo calcolato": 0, "Selezione": False},
    ]
)

SECTION_OPTIONS = {
    "Materialita preliminare": "preliminare",
    "Materialita definitiva": "definitiva",
}


def _to_int_series(series, default_value):
    return pd.to_numeric(series, errors="coerce").fillna(default_value).round(0).astype(int)


def _load_basi_per_anno(conn):
    return pd.read_sql(
        """
        WITH latest_schema AS (
            SELECT MAX(id) AS id FROM lead_schema_version
        ),
        latest_mapping AS (
            SELECT *
            FROM (
                SELECT m.*,
                       ROW_NUMBER() OVER (
                           PARTITION BY m.gl_account_id
                           ORDER BY m.schema_version_id DESC, m.id DESC
                       ) AS rn
                FROM account_lead_mapping m
                WHERE m.is_active = 1
            )
            WHERE rn = 1
        ),
        base_data AS (
            SELECT
                tbh.fiscal_year,
                UPPER(TRIM(COALESCE(ls.tipo, ''))) AS tipo_norm,
                UPPER(TRIM(COALESCE(ls.lead, ''))) AS lead_norm,
                UPPER(TRIM(COALESCE(ls.sublead, ''))) AS sublead_norm,
                tbl.closing_balance AS importo
            FROM latest_mapping m
            JOIN latest_schema s ON 1 = 1
            JOIN lead_structure ls
              ON ls.schema_version_id = s.id
             AND ls.sublead = m.sublead
            JOIN trial_balance_line tbl ON tbl.gl_account_id = m.gl_account_id
            JOIN trial_balance_header tbh ON tbh.id = tbl.trial_balance_id
        )
        SELECT
            fiscal_year,
            SUM(CASE WHEN sublead_norm = 'U0100' THEN importo ELSE 0 END) AS ricavi_u0100,
            SUM(CASE WHEN tipo_norm = 'ATTIVO' THEN importo ELSE 0 END) AS totale_attivo,
            SUM(CASE WHEN lead_norm = 'L PATRIMONIO NETTO' THEN importo ELSE 0 END) AS patrimonio_netto_exact,
            SUM(CASE WHEN lead_norm = 'L' THEN importo ELSE 0 END) AS patrimonio_netto_fallback,
            SUM(CASE WHEN tipo_norm = 'CE' AND lead_norm <> 'YF' THEN importo ELSE 0 END) AS reddito_ante_imposte
        FROM base_data
        GROUP BY fiscal_year
        ORDER BY fiscal_year DESC
        """,
        conn,
    )


def _format_int_it(value):
    return f"{int(abs(round(value, 0))):,}".replace(",", ".")


def _build_excel_export(df_export, df_summary):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df_export.to_excel(writer, index=False, sheet_name="Criteri")
        df_summary.to_excel(writer, index=False, sheet_name="Riepilogo")
    return output.getvalue()


def _build_word_export(df_export, df_summary):
    try:
        from docx import Document
        import io as _io
    except ImportError:
        return None

    doc = Document()
    doc.add_heading("Determinazione Materialita", level=1)

    doc.add_heading("Riepilogo", level=2)
    for _, row in df_summary.iterrows():
        doc.add_paragraph(f"{row['Voce']}: {row['Valore']}")

    doc.add_heading("Criteri", level=2)
    table = doc.add_table(rows=1, cols=len(df_export.columns))
    table.style = "Table Grid"
    for i, c in enumerate(df_export.columns):
        table.rows[0].cells[i].text = str(c)
    for _, row in df_export.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(df_export.columns):
            cells[i].text = str(row[c])

    buffer = _io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _build_pdf_export(df_export, df_summary):
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet
        import io as _io
    except ImportError:
        return None

    buffer = _io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(A4))
    styles = getSampleStyleSheet()
    elements = [Paragraph("Determinazione Materialita", styles["Heading2"]), Spacer(1, 8)]

    elements.append(Paragraph("Riepilogo", styles["Heading3"]))
    summary_data = [["Voce", "Valore"]]
    for _, row in df_summary.iterrows():
        voce_text = Paragraph(str(row["Voce"]), styles["BodyText"])
        valore_text = Paragraph(str(row["Valore"]), styles["BodyText"])
        summary_data.append([voce_text, valore_text])
    summary_table = Table(summary_data, colWidths=[220, 420])
    summary_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EEEEEE")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    elements.append(summary_table)
    elements.append(Spacer(1, 10))

    elements.append(Paragraph("Criteri", styles["Heading3"]))
    criteria_data = [df_export.columns.tolist()] + df_export.values.tolist()
    criteria_table = Table(criteria_data, repeatRows=1)
    criteria_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EEEEEE")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
    ]))
    elements.append(criteria_table)

    doc.build(elements)
    return buffer.getvalue()


try:
    conn = get_conn()
    df_basi = _load_basi_per_anno(conn)

    if df_basi.empty:
        st.warning("Nessun valore disponibile per i criteri di materialita.")
    else:
        with st.sidebar.expander("📘 Determinazione della Materialita – Riferimenti ISA", expanded=False):
            st.markdown(
                """
                **Materialita per il bilancio nel suo complesso**
                
                Ai sensi dell'ISA 320, par. 10, il revisore determina la materialita nella fase di pianificazione, esercitando il proprio giudizio professionale.  
                La determinazione richiede:
                - individuazione di una base di riferimento appropriata (ISA 320 A3-A7);
                - applicazione di una percentuale coerente con la prassi professionale;
                - considerazione di fattori qualitativi rilevanti per gli utilizzatori del bilancio.

                Ove appropriato, possono essere determinate materialita specifiche per particolari classi di operazioni o informativa (ISA 320 A10-A12).

                **Materialita operativa**

                Conformemente all'ISA 320, par. 9, la materialita operativa e fissata a un livello inferiore rispetto alla materialita complessiva, al fine di ridurre a un livello appropriatamente basso il rischio che errori non corretti e non individuati eccedano la materialita per il bilancio nel suo complesso (ISA 320 A13-A14).

                La sua determinazione tiene conto della valutazione dei rischi ai sensi dell'ISA 315 e dei risultati di revisioni precedenti.

                **Errore chiaramente trascurabile**

                Ai sensi dell'ISA 450, par. 5 e A2-A3, il revisore puo stabilire una soglia al di sotto della quale gli errori sono considerati chiaramente trascurabili e non soggetti ad accumulo.

                Tale soglia e significativamente inferiore alla materialita e non sostituisce la necessaria valutazione qualitativa degli errori.

                **Documentazione**

                Le basi adottate, le percentuali applicate e le motivazioni devono essere adeguatamente documentate in conformita all'ISA 230.
                """
            )
        st.sidebar.subheader("Materialita")
        selected_section_label = st.sidebar.radio(
            "Sezione",
            options=list(SECTION_OPTIONS.keys()),
            index=0,
        )
        section_key = SECTION_OPTIONS[selected_section_label]
        st.subheader(selected_section_label)

        fiscal_years = [int(y) for y in df_basi["fiscal_year"].dropna().astype(int).tolist()]
        default_index = 1 if section_key == "preliminare" and len(fiscal_years) > 1 else 0
        selected_year = st.selectbox(
            "Bilancio da considerare",
            options=fiscal_years,
            index=default_index,
            key=f"{section_key}_bilancio_anno",
        )

        df_selected = df_basi[df_basi["fiscal_year"].astype(int) == int(selected_year)]
        if df_selected.empty:
            st.warning("Nessun dato disponibile per l'esercizio selezionato.")
            st.stop()

        row = df_selected.iloc[0]
        patrimonio_netto_exact = float(row["patrimonio_netto_exact"] or 0)
        patrimonio_netto_fallback = float(row["patrimonio_netto_fallback"] or 0)
        patrimonio_netto = patrimonio_netto_exact if patrimonio_netto_exact != 0 else patrimonio_netto_fallback
        basi_map = {
            "A1) Ricavi delle vendite e delle prestazioni": float(row["ricavi_u0100"] or 0),
            "Totale attivo": float(row["totale_attivo"] or 0),
            "Patrimonio netto": patrimonio_netto,
            "Reddito ante imposte": float(row["reddito_ante_imposte"] or 0),
        }

        editable_key = f"materialita_editable_{section_key}"
        if editable_key not in st.session_state:
            st.session_state[editable_key] = DEFAULT_ROWS[["Voce", "% selezionata", "Selezione"]].copy()

        table_df = DEFAULT_ROWS.copy()
        editable_prev = st.session_state[editable_key].copy()
        if not editable_prev.empty and "Voce" in editable_prev.columns:
            editable_prev = editable_prev.set_index("Voce")
            for col in ["% selezionata", "Selezione"]:
                if col in editable_prev.columns:
                    table_df[col] = table_df["Voce"].map(editable_prev[col]).fillna(table_df[col])

        table_df["% min"] = _to_int_series(table_df["% min"], 0)
        table_df["% max"] = _to_int_series(table_df["% max"], 100)
        table_df["% selezionata"] = pd.to_numeric(table_df["% selezionata"], errors="coerce")
        table_df["% selezionata"] = table_df["% selezionata"].fillna(table_df["% min"])
        table_df["% selezionata"] = table_df["% selezionata"].clip(lower=table_df["% min"], upper=table_df["% max"]).round(0).astype(int)
        table_df["Valore base"] = table_df["Voce"].map(basi_map).fillna(0.0).abs().round(0).astype(int)
        table_df["Importo calcolato"] = (table_df["Valore base"] * table_df["% selezionata"] / 100).round(0).astype(int)
        table_df = table_df[["Voce", "Valore base", "% min", "% max", "% selezionata", "Importo calcolato", "Selezione"]]

        edited = st.data_editor(
            table_df,
            use_container_width=True,
            hide_index=True,
            row_height=24,
            column_config={
                "Voce": st.column_config.TextColumn("Voce"),
                "Valore base": st.column_config.NumberColumn("Valore base", format="%d"),
                "% min": st.column_config.NumberColumn("% min", min_value=0, max_value=100, step=1, format="%d"),
                "% max": st.column_config.NumberColumn("% max", min_value=0, max_value=100, step=1, format="%d"),
                "% selezionata": st.column_config.SelectboxColumn("% selezionata", options=[1, 2, 3, 4, 5, 6, 7], required=True),
                "Importo calcolato": st.column_config.NumberColumn("Importo calcolato", format="%d"),
                "Selezione": st.column_config.CheckboxColumn("Selezione"),
            },
            disabled=["Voce", "% min", "% max", "Valore base", "Importo calcolato"],
        )

        edited["% min"] = _to_int_series(edited["% min"], 0)
        edited["% max"] = _to_int_series(edited["% max"], 100)
        edited["% selezionata"] = pd.to_numeric(edited["% selezionata"], errors="coerce")
        edited["% selezionata"] = edited["% selezionata"].fillna(edited["% min"])
        edited["% selezionata"] = edited["% selezionata"].clip(lower=edited["% min"], upper=edited["% max"]).round(0).astype(int)
        edited["Selezione"] = edited["Selezione"].fillna(False).astype(bool)
        edited["Valore base"] = edited["Voce"].map(basi_map).fillna(0.0).abs().round(0).astype(int)
        edited["Importo calcolato"] = (edited["Valore base"] * edited["% selezionata"] / 100).round(0).astype(int)

        new_editable = edited[["Voce", "% selezionata", "Selezione"]].copy()
        prev_editable = st.session_state[editable_key].copy()
        st.session_state[editable_key] = new_editable
        if not new_editable.equals(prev_editable):
            st.rerun()

        media_materialita = None
        materialita_operativa = None
        errori_trascurabili = None
        pct_materialita_operativa = None
        pct_errori_trascurabili = None
        note_default_text = (
            "Inserire una descrizione del criterio selezionato per la determinazione della materialita, "
            "specificando le ragioni professionali della scelta delle percentuali applicate ai benchmark "
            "considerati e gli elementi qualitativi/quantitativi rilevanti emersi nell'analisi."
        )
        note_state_key = f"nota_materialita_store_{section_key}"
        note_widget_key = f"nota_materialita_input_{section_key}"
        if note_state_key not in st.session_state:
            st.session_state[note_state_key] = note_default_text
        if note_widget_key not in st.session_state:
            st.session_state[note_widget_key] = st.session_state[note_state_key]

        importi_selezionati = edited.loc[edited["Selezione"], "Importo calcolato"]
        if importi_selezionati.empty:
            st.warning("E' necessario selezionare almeno un criterio di determinazione")
        else:
            media_materialita = int(round(importi_selezionati.mean(), 0))
            st.metric("Materialita generale", _format_int_it(media_materialita))

            col_mo_slider, col_mo_metric = st.columns([1, 2.2])
            with col_mo_slider:
                pct_materialita_operativa = st.slider(
                    "% Materialita operativa",
                    min_value=60,
                    max_value=80,
                    value=75,
                    step=1,
                    format="%d%%",
                    key=f"slider_materialita_operativa_{section_key}",
                )
            materialita_operativa = int(round(media_materialita * pct_materialita_operativa / 100, 0))
            with col_mo_metric:
                st.metric("Materialita operativa", _format_int_it(materialita_operativa))

            col_et_slider, col_et_metric = st.columns([1, 2.2])
            with col_et_slider:
                pct_errori_trascurabili = st.slider(
                    "% Errori trascurabili",
                    min_value=5,
                    max_value=15,
                    value=10,
                    step=1,
                    format="%d%%",
                    key=f"slider_errori_trascurabili_{section_key}",
                )
            errori_trascurabili = int(round(materialita_operativa * pct_errori_trascurabili / 100, 0))
            with col_et_metric:
                st.metric("Errori trascurabili", _format_int_it(errori_trascurabili))

        nota_text = st.text_area(
            "Spiegazione del criterio utilizzato e relative motivazioni",
            height=140,
            key=note_widget_key,
        )
        st.session_state[note_state_key] = nota_text

        df_export = edited.copy()
        df_export["Selezione"] = df_export["Selezione"].map(lambda x: "Si" if bool(x) else "No")
        for col in ["Valore base", "Importo calcolato"]:
            df_export[col] = df_export[col].astype(int)

        df_summary = pd.DataFrame(
            [
                {"Voce": "Sezione", "Valore": selected_section_label},
                {"Voce": "Esercizio selezionato", "Valore": int(selected_year)},
                {"Voce": "Materialita generale", "Valore": _format_int_it(media_materialita) if media_materialita is not None else ""},
                {"Voce": "% Materialita operativa", "Valore": f"{pct_materialita_operativa}%" if pct_materialita_operativa is not None else ""},
                {"Voce": "Materialita operativa", "Valore": _format_int_it(materialita_operativa) if materialita_operativa is not None else ""},
                {"Voce": "% Errori trascurabili", "Valore": f"{pct_errori_trascurabili}%" if pct_errori_trascurabili is not None else ""},
                {"Voce": "Errori trascurabili", "Valore": _format_int_it(errori_trascurabili) if errori_trascurabili is not None else ""},
                {"Voce": "Nota", "Valore": nota_text},
            ]
        )

        excel_data = _build_excel_export(df_export, df_summary)
        word_data = _build_word_export(df_export, df_summary)
        pdf_data = _build_pdf_export(df_export, df_summary)

        col_ex, col_wd, col_pdf = st.columns(3)
        with col_ex:
            st.download_button(
                label="Esporta in Excel",
                data=excel_data,
                file_name=f"materialita_{section_key}_{selected_year}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        with col_wd:
            if word_data is None:
                st.warning("Export Word non disponibile: installare `python-docx`.")
            else:
                st.download_button(
                    label="Esporta in Word",
                    data=word_data,
                    file_name=f"materialita_{section_key}_{selected_year}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
        with col_pdf:
            if pdf_data is None:
                st.warning("Export PDF non disponibile: installare `reportlab`.")
            else:
                st.download_button(
                    label="Esporta in PDF",
                    data=pdf_data,
                    file_name=f"materialita_{section_key}_{selected_year}.pdf",
                    mime="application/pdf",
                )

        st.caption(
            " | ".join(
                [
                    f"Esercizio selezionato: {int(selected_year)}",
                    f"Ricavi: {_format_int_it(basi_map['A1) Ricavi delle vendite e delle prestazioni'])}",
                    f"Totale attivo: {_format_int_it(basi_map['Totale attivo'])}",
                    f"Patrimonio netto: {_format_int_it(basi_map['Patrimonio netto'])}",
                    f"Reddito ante imposte: {_format_int_it(basi_map['Reddito ante imposte'])}",
                ]
            )
        )

except Exception as e:
    st.error("Errore nella pagina Materialita.")
    st.exception(e)

finally:
    try:
        conn.close()
    except Exception:
        pass
