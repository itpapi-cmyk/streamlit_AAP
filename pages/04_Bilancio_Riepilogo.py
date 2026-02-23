import streamlit as st
import pandas as pd
from modules.lead_numeric.ddl import init_db
from modules.lead_numeric.db import get_conn

st.set_page_config(page_title="04 — Bilancio Riepilogo", layout="wide")
st.title("04 — Bilancio: Lead, Conto, Importo")

init_db()
TIPO_ORDER = ["ATTIVO", "PASSIVO", "CE"]
TIPO_LABELS = {"ATTIVO": "ATTIVO", "PASSIVO": "PASSIVO", "CE": "CONTO ECONOMICO"}
VIEW_OPTIONS = {
    "lead_dettaglio": "Lead dettaglio",
    "lead": "Lead",
    "subtotali_lead": "Subtotali Lead",
    "gruppo_lead": "Gruppo Lead",
    "totali_tipo": "Totali per tipo",
}
TABLE_ROW_HEIGHT = 24


def _format_percent_it(value):
    if pd.isna(value):
        return ""
    return f"{value:,.2f}%".replace(",", "X").replace(".", ",").replace("X", ".")


def _scale_amount_columns(dataframe, amount_cols, scale_factor):
    if scale_factor == 1:
        return dataframe
    scaled_df = dataframe.copy()
    for col in amount_cols:
        if col in scaled_df.columns:
            scaled_df[col] = pd.to_numeric(scaled_df[col], errors="coerce") / scale_factor
    return scaled_df


def _style_bilancio_table(dataframe, amount_cols, amount_decimals=2):
    number_fmt = "{:,.%df}" % amount_decimals
    format_map = {col: number_fmt for col in amount_cols}
    format_map["differenza_percentuale"] = _format_percent_it
    def _row_style(row):
        label = str(row.get("account_name", "") or row.get("descr_sublead", ""))
        tipo_value = str(row.get("tipo", "")).strip().upper()
        if tipo_value == "CHECK":
            return ["font-weight: 800; background-color: #F2F2F2; border-top: 2px solid #9E9E9E;"] * len(row)
        if label.startswith("Totale LEAD "):
            return ["font-weight: 700; background-color: #E8F4FF; border-top: 1px solid #7DB5E8;"] * len(row)
        if label.startswith("Totale GROUP_LEAD "):
            return ["font-weight: 700; background-color: #FFF4E5; border-top: 1px solid #E2A35A;"] * len(row)
        if label.startswith("Totale TIPO "):
            return ["font-weight: 800; background-color: #EAF9EA; border-top: 2px solid #63A35C; border-bottom: 2px solid #63A35C;"] * len(row)
        return [""] * len(row)

    return dataframe.style.format(format_map, decimal=",", thousands=".").apply(_row_style, axis=1)


def _render_bilancio_dataframe(styled_df):
    st.dataframe(
        styled_df,
        use_container_width=True,
        hide_index=True,
        row_height=TABLE_ROW_HEIGHT,
    )


def _append_check_row(df_totali_tipo, latest_col, previous_col):
    prev_total = df_totali_tipo[previous_col].fillna(0).sum()
    latest_total = df_totali_tipo[latest_col].fillna(0).sum()
    diff_total = df_totali_tipo["differenza_valore"].fillna(0).sum()
    pct_total = (diff_total / prev_total * 100) if prev_total != 0 else None

    check_row = {c: "" for c in df_totali_tipo.columns}
    check_row["tipo"] = "check"
    check_row[latest_col] = latest_total
    check_row[previous_col] = prev_total
    check_row["differenza_valore"] = diff_total
    check_row["differenza_percentuale"] = pct_total
    return pd.concat([df_totali_tipo, pd.DataFrame([check_row])], ignore_index=True)


def _format_number_it(value, scale_factor=1, decimals=2):
    if pd.isna(value):
        return ""
    scaled_value = value / scale_factor if scale_factor else value
    return f"{scaled_value:,.{decimals}f}".replace(",", "X").replace(".", ",").replace("X", ".")


def _build_pdf_by_lead(df_source, latest_col, previous_col, latest_year, previous_year, amount_scale=1, amount_decimals=2):
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, PageBreak
        import io
    except ImportError:
        return None

    pdf_cols = [
        "lead", "sublead", "descr_sublead", "account_code", "account_name",
        latest_col, previous_col, "differenza_valore", "differenza_percentuale"
    ]
    missing_cols = [c for c in pdf_cols if c not in df_source.columns]
    if missing_cols:
        return None

    df_pdf = (
        df_source[pdf_cols]
        .copy()
        .sort_values(["lead", "sublead", "account_code"], kind="stable")
    )

    headers = [
        "Account Code",
        "Account Name",
        f"Importo {latest_year}",
        f"Importo {previous_year}",
        "Diff valore",
        "Diff %",
    ]

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        leftMargin=18,
        rightMargin=18,
        topMargin=18,
        bottomMargin=18,
    )
    styles = getSampleStyleSheet()
    elements = []

    for idx, (lead_value, df_lead) in enumerate(df_pdf.groupby("lead", sort=False)):
        if idx > 0:
            elements.append(PageBreak())

        elements.append(Paragraph(f"Lead: {lead_value}", styles["Heading3"]))
        elements.append(
            Paragraph(
                f"Confronto periodi: {latest_year} vs {previous_year}",
                styles["Normal"]
            )
        )
        elements.append(Spacer(1, 6))

        for sublead_value, df_sublead in df_lead.groupby("sublead", sort=False):
            descr_value = str(df_sublead["descr_sublead"].iloc[0] or "")
            elements.append(
                Paragraph(
                    f"Sublead: {sublead_value} - {descr_value}",
                    styles["Heading4"]
                )
            )
            elements.append(Spacer(1, 4))

            rows = [headers]
            for _, row in df_sublead.iterrows():
                rows.append([
                    str(row.get("account_code", "") or ""),
                    str(row.get("account_name", "") or ""),
                    _format_number_it(row.get(latest_col), amount_scale, amount_decimals),
                    _format_number_it(row.get(previous_col), amount_scale, amount_decimals),
                    _format_number_it(row.get("differenza_valore"), amount_scale, amount_decimals),
                    _format_percent_it(row.get("differenza_percentuale")),
                ])

            prev_sublead_total = df_sublead[previous_col].fillna(0).sum()
            latest_sublead_total = df_sublead[latest_col].fillna(0).sum()
            diff_sublead_total = latest_sublead_total - prev_sublead_total
            pct_sublead_total = (diff_sublead_total / prev_sublead_total * 100) if prev_sublead_total != 0 else None
            rows.append([
                "",
                "Totale Sublead",
                _format_number_it(latest_sublead_total, amount_scale, amount_decimals),
                _format_number_it(prev_sublead_total, amount_scale, amount_decimals),
                _format_number_it(diff_sublead_total, amount_scale, amount_decimals),
                _format_percent_it(pct_sublead_total),
            ])

            table = Table(
                rows,
                repeatRows=1,
                colWidths=[80, 260, 95, 95, 95, 65],
            )
            table.setStyle(TableStyle([
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -2), "Helvetica"),
                ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EDEDED")),
                ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#F5F5F5")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#BDBDBD")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (2, 1), (5, -1), "RIGHT"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]))
            elements.append(table)
            elements.append(Spacer(1, 8))

        prev_lead_total = df_lead[previous_col].fillna(0).sum()
        latest_lead_total = df_lead[latest_col].fillna(0).sum()
        diff_lead_total = latest_lead_total - prev_lead_total
        pct_lead_total = (diff_lead_total / prev_lead_total * 100) if prev_lead_total != 0 else None

        lead_total_table = Table(
            [[
                "Totale Lead",
                _format_number_it(latest_lead_total, amount_scale, amount_decimals),
                _format_number_it(prev_lead_total, amount_scale, amount_decimals),
                _format_number_it(diff_lead_total, amount_scale, amount_decimals),
                _format_percent_it(pct_lead_total),
            ]],
            colWidths=[340, 95, 95, 95, 65],
        )
        lead_total_table.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#E8F4FF")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#7DB5E8")),
            ("ALIGN", (1, 0), (4, 0), "RIGHT"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        elements.append(lead_total_table)

    doc.build(elements)
    return buffer.getvalue()


def _build_docx_by_lead(df_source, latest_col, previous_col, latest_year, previous_year, amount_scale=1, amount_decimals=2):
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        import io
    except ImportError:
        return None

    word_cols = [
        "lead", "sublead", "descr_sublead", "account_code", "account_name",
        latest_col, previous_col, "differenza_valore", "differenza_percentuale"
    ]
    missing_cols = [c for c in word_cols if c not in df_source.columns]
    if missing_cols:
        return None

    df_word = (
        df_source[word_cols]
        .copy()
        .sort_values(["lead", "sublead", "account_code"], kind="stable")
    )

    headers = [
        "Account Code",
        "Account Name",
        f"Importo {latest_year}",
        f"Importo {previous_year}",
        "Diff valore",
        "Diff %",
    ]

    doc = Document()

    for idx, (lead_value, df_lead) in enumerate(df_word.groupby("lead", sort=False)):
        if idx > 0:
            doc.add_page_break()

        doc.add_heading(f"Lead: {lead_value}", level=2)
        p_confronto = doc.add_paragraph(f"Confronto periodi: {latest_year} vs {previous_year}")
        p_confronto.runs[0].font.size = Pt(10)

        for sublead_value, df_sublead in df_lead.groupby("sublead", sort=False):
            descr_value = str(df_sublead["descr_sublead"].iloc[0] or "")
            doc.add_heading(f"Sublead: {sublead_value} - {descr_value}", level=3)

            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for col_idx, header in enumerate(headers):
                run = hdr_cells[col_idx].paragraphs[0].add_run(header)
                run.bold = True
                if col_idx >= 2:
                    hdr_cells[col_idx].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            for _, row in df_sublead.iterrows():
                data_cells = table.add_row().cells
                data_cells[0].text = str(row.get("account_code", "") or "")
                data_cells[1].text = str(row.get("account_name", "") or "")
                data_cells[2].text = _format_number_it(row.get(latest_col), amount_scale, amount_decimals)
                data_cells[3].text = _format_number_it(row.get(previous_col), amount_scale, amount_decimals)
                data_cells[4].text = _format_number_it(row.get("differenza_valore"), amount_scale, amount_decimals)
                data_cells[5].text = _format_percent_it(row.get("differenza_percentuale"))
                for col_idx in [2, 3, 4, 5]:
                    data_cells[col_idx].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            prev_sublead_total = df_sublead[previous_col].fillna(0).sum()
            latest_sublead_total = df_sublead[latest_col].fillna(0).sum()
            diff_sublead_total = latest_sublead_total - prev_sublead_total
            pct_sublead_total = (diff_sublead_total / prev_sublead_total * 100) if prev_sublead_total != 0 else None

            total_cells = table.add_row().cells
            total_cells[0].text = ""
            total_cells[1].text = "Totale Sublead"
            total_cells[2].text = _format_number_it(latest_sublead_total, amount_scale, amount_decimals)
            total_cells[3].text = _format_number_it(prev_sublead_total, amount_scale, amount_decimals)
            total_cells[4].text = _format_number_it(diff_sublead_total, amount_scale, amount_decimals)
            total_cells[5].text = _format_percent_it(pct_sublead_total)
            for cell in total_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            for col_idx in [2, 3, 4, 5]:
                total_cells[col_idx].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            doc.add_paragraph("")

        prev_lead_total = df_lead[previous_col].fillna(0).sum()
        latest_lead_total = df_lead[latest_col].fillna(0).sum()
        diff_lead_total = latest_lead_total - prev_lead_total
        pct_lead_total = (diff_lead_total / prev_lead_total * 100) if prev_lead_total != 0 else None

        p_total_lead = doc.add_paragraph(
            "Totale Lead | "
            f"Importo {latest_year}: {_format_number_it(latest_lead_total, amount_scale, amount_decimals)} | "
            f"Importo {previous_year}: {_format_number_it(prev_lead_total, amount_scale, amount_decimals)} | "
            f"Diff valore: {_format_number_it(diff_lead_total, amount_scale, amount_decimals)} | "
            f"Diff %: {_format_percent_it(pct_lead_total)}"
        )
        for run in p_total_lead.runs:
            run.bold = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def _subtotal_row(base_cols, latest_col, previous_col, tipo, group_lead, lead, label, source_df, label_col):
    prev_total = source_df[previous_col].fillna(0).sum()
    latest_total = source_df[latest_col].fillna(0).sum()
    diff_total = latest_total - prev_total
    pct_total = (diff_total / prev_total * 100) if prev_total != 0 else None

    row = {c: "" for c in base_cols}
    row["tipo"] = tipo
    row["group_lead"] = group_lead
    row["lead"] = lead
    row[label_col] = label
    row[latest_col] = latest_total
    row[previous_col] = prev_total
    row["differenza_valore"] = diff_total
    row["differenza_percentuale"] = pct_total
    return row

def _build_subtotals(df_source, group_col, latest_col, previous_col):
    subtot = (
        df_source
        .groupby(group_col, as_index=False)[[latest_col, previous_col]]
        .sum()
    )
    subtot["differenza_valore"] = subtot[latest_col].fillna(0) - subtot[previous_col].fillna(0)
    subtot["differenza_percentuale"] = subtot.apply(
        lambda r: (r["differenza_valore"] / r[previous_col] * 100)
        if pd.notnull(r[previous_col]) and r[previous_col] != 0 else None,
        axis=1
    )
    return subtot


def _prepare_export_dataframe(dataframe, amount_cols, scale_factor, decimals):
    df_export = _scale_amount_columns(dataframe, amount_cols, scale_factor)
    if decimals is not None:
        for col in amount_cols:
            if col in df_export.columns:
                numeric_col = pd.to_numeric(df_export[col], errors="coerce")
                df_export[col] = numeric_col.round(decimals)
    return df_export

def _build_bilancio_with_break_subtotals(df_pivot, ordered_cols, latest_col, previous_col, label_col):
    base_cols = [c for c in ordered_cols if c not in [latest_col, previous_col, "differenza_valore", "differenza_percentuale"]]
    rows = []

    sort_cols = [c for c in ["tipo_subtotale", "group_lead", "lead", "account_code", "sublead"] if c in df_pivot.columns]
    df_sorted = df_pivot.sort_values(sort_cols, kind="stable").copy()
    df_sorted["tipo_subtotale"] = pd.Categorical(
        df_sorted["tipo_subtotale"],
        categories=TIPO_ORDER,
        ordered=True
    )

    for tipo_value in TIPO_ORDER:
        df_tipo = df_sorted[df_sorted["tipo_subtotale"] == tipo_value]
        if df_tipo.empty:
            rows.append(
                _subtotal_row(
                    ordered_cols, latest_col, previous_col,
                    tipo_value, "", "",
                    f"Totale TIPO {tipo_value}", df_tipo, label_col
                )
            )
            continue
        for group_value, df_group in df_tipo.groupby("group_lead", sort=False):
            for lead_value, df_lead in df_group.groupby("lead", sort=False):
                for _, detail_row in df_lead.iterrows():
                    rows.append(detail_row[ordered_cols].to_dict())
                rows.append(
                    _subtotal_row(
                        ordered_cols, latest_col, previous_col,
                        tipo_value, group_value, lead_value,
                        f"Totale LEAD {lead_value}", df_lead, label_col
                    )
                )

            rows.append(
                _subtotal_row(
                    ordered_cols, latest_col, previous_col,
                    tipo_value, group_value, "",
                    f"Totale GROUP_LEAD {group_value}", df_group, label_col
                )
            )

        rows.append(
            _subtotal_row(
                ordered_cols, latest_col, previous_col,
                tipo_value, "", "",
                f"Totale TIPO {tipo_value}", df_tipo, label_col
            )
        )

    return pd.DataFrame(rows, columns=ordered_cols)


try:
    conn = get_conn()

    # Recupera dati di bilancio con mapping
    # Recupera dati per 2024 e 2025
    df = pd.read_sql(
        '''
        WITH latest_schema AS (
            SELECT MAX(id) AS id FROM lead_schema_version
        ),
        latest_mapping AS (
            SELECT *
            FROM (
                SELECT m.*,
                       ROW_NUMBER() OVER (
                           PARTITION BY m.gl_account_id
                           ORDER BY m.schema_version_id DESC, m.id DESC
                       ) AS rn
                FROM account_lead_mapping m
                WHERE m.is_active = 1
            )
            WHERE rn = 1
        )
        SELECT ls.group_lead, ls.tipo, ls.lead, ls.sublead, ls.descrizione_cee AS descr_sublead, ga.account_code, ga.account_name,
               tbl.closing_balance AS importo, tbh.fiscal_year
        FROM latest_mapping m
        JOIN gl_account ga ON ga.id = m.gl_account_id
        JOIN latest_schema s ON 1=1
        JOIN lead_structure ls ON ls.sublead = m.sublead AND ls.schema_version_id = s.id
        LEFT JOIN trial_balance_line tbl ON tbl.gl_account_id = ga.id
        LEFT JOIN trial_balance_header tbh ON tbl.trial_balance_id = tbh.id
        WHERE tbh.fiscal_year IS NOT NULL
        ORDER BY ls.group_lead, ls.tipo, ls.lead, ga.account_code, tbh.fiscal_year
        ''', conn)

    if df.empty:
        st.info("Nessun dato di bilancio disponibile.")
    else:
        years = sorted(df['fiscal_year'].dropna().astype(int).unique().tolist(), reverse=True)
        if len(years) < 2:
            st.warning("Servono almeno due esercizi per calcolare il confronto.")
            st.stop()

        st.subheader("Selezione periodi di confronto")
        col_year_1, col_year_2 = st.columns(2)

        with col_year_1:
            latest_year = st.selectbox(
                "Periodo 1",
                options=years,
                index=0,
                help="Anno principale del confronto.",
            )

        previous_options = [y for y in years if y != latest_year]
        default_previous = years[1] if len(years) > 1 else years[0]
        default_previous_index = previous_options.index(default_previous) if default_previous in previous_options else 0
        with col_year_2:
            previous_year = st.selectbox(
                "Periodo 2",
                options=previous_options,
                index=default_previous_index,
                help="Anno di confronto.",
            )

        df = df[df["fiscal_year"].astype(int).isin([latest_year, previous_year])].copy()
        latest_col = f'importo_{latest_year}'
        previous_col = f'importo_{previous_year}'
        index_cols = ['tipo', 'group_lead', 'lead', 'sublead', 'descr_sublead', 'account_code', 'account_name']

        # Pivot per confronto anno più recente vs anno precedente
        df_pivot = df.pivot_table(
            index=index_cols,
            columns='fiscal_year',
            values='importo',
            aggfunc='sum'
        ).reset_index()
        df_pivot.columns.name = None
        df_pivot = df_pivot.rename(columns={year: f'importo_{int(year)}' for year in years})
        for col in [latest_col, previous_col]:
            if col not in df_pivot.columns:
                df_pivot[col] = 0
        df_pivot['differenza_valore'] = df_pivot[latest_col].fillna(0) - df_pivot[previous_col].fillna(0)
        df_pivot['differenza_percentuale'] = df_pivot.apply(
            lambda r: (r['differenza_valore'] / r[previous_col] * 100) if pd.notnull(r[previous_col]) and r[previous_col] != 0 else None,
            axis=1
        )
        tipo_valid_mask = df_pivot['tipo'].notna() & (df_pivot['tipo'].astype(str).str.strip() != "")
        df_pivot['tipo_subtotale'] = (
            df_pivot['tipo']
            .where(tipo_valid_mask, 'CE')
            .astype(str)
            .str.strip()
            .str.upper()
        )
        ordered_cols = index_cols + [latest_col, previous_col, 'differenza_valore', 'differenza_percentuale']
        df_display = _build_bilancio_with_break_subtotals(
            df_pivot=df_pivot,
            ordered_cols=ordered_cols,
            latest_col=latest_col,
            previous_col=previous_col,
            label_col="account_name"
        )
        index_cols_no_account = ['tipo', 'group_lead', 'lead', 'sublead', 'descr_sublead']
        df_no_account = (
            df_pivot
            .groupby(index_cols_no_account, as_index=False)[[latest_col, previous_col]]
            .sum()
        )
        df_no_account['differenza_valore'] = df_no_account[latest_col].fillna(0) - df_no_account[previous_col].fillna(0)
        df_no_account['differenza_percentuale'] = df_no_account.apply(
            lambda r: (r['differenza_valore'] / r[previous_col] * 100) if pd.notnull(r[previous_col]) and r[previous_col] != 0 else None,
            axis=1
        )
        tipo_valid_mask_2 = df_no_account['tipo'].notna() & (df_no_account['tipo'].astype(str).str.strip() != "")
        df_no_account['tipo_subtotale'] = (
            df_no_account['tipo']
            .where(tipo_valid_mask_2, 'CE')
            .astype(str)
            .str.strip()
            .str.upper()
        )
        ordered_cols_no_account = index_cols_no_account + [latest_col, previous_col, 'differenza_valore', 'differenza_percentuale']
        df_display_no_account = _build_bilancio_with_break_subtotals(
            df_pivot=df_no_account,
            ordered_cols=ordered_cols_no_account,
            latest_col=latest_col,
            previous_col=previous_col,
            label_col="descr_sublead"
        )
        subtot_lead = _build_subtotals(df_pivot, 'lead', latest_col, previous_col)
        subtot_group_lead = _build_subtotals(df_pivot, 'group_lead', latest_col, previous_col)
        subtot_tipo = _build_subtotals(df_pivot, 'tipo_subtotale', latest_col, previous_col)
        subtot_tipo = subtot_tipo.rename(columns={'tipo_subtotale': 'tipo'})
        subtot_tipo = subtot_tipo.set_index("tipo").reindex(TIPO_ORDER, fill_value=0).reset_index()
        subtot_tipo["tipo"] = subtot_tipo["tipo"].map(TIPO_LABELS).fillna(subtot_tipo["tipo"])
        subtot_tipo = _append_check_row(subtot_tipo, latest_col, previous_col)

        amount_unit = st.sidebar.radio(
            "Unità importi",
            options=["euro", "euro_1000"],
            format_func=lambda x: "Euro" if x == "euro" else "Euro/1000 (1 decimale)",
            index=0
        )
        amount_scale = 1000 if amount_unit == "euro_1000" else 1
        amount_decimals = 1 if amount_unit == "euro_1000" else 2
        amount_cols = [latest_col, previous_col, "differenza_valore"]

        selected_view = st.sidebar.radio(
            "Vista Bilancio Riepilogo",
            options=list(VIEW_OPTIONS.keys()),
            format_func=lambda x: VIEW_OPTIONS[x],
            index=0
        )

        # Subtotali per Lead, Group Lead, Tipo (anno più recente/precedente e differenze)
        if selected_view == "lead_dettaglio":
            st.subheader("Lead dettaglio")
            _render_bilancio_dataframe(
                _style_bilancio_table(
                    _scale_amount_columns(df_display, amount_cols, amount_scale),
                    amount_cols,
                    amount_decimals
                )
            )
            st.caption(
                f"Riepilogo: Lead, Conto COGE, Importo {latest_year}, Importo {previous_year}, Differenza valore e %."
            )
        elif selected_view == "lead":
            st.subheader("Lead")
            _render_bilancio_dataframe(
                _style_bilancio_table(
                    _scale_amount_columns(df_display_no_account, amount_cols, amount_scale),
                    amount_cols,
                    amount_decimals
                )
            )
        elif selected_view == "subtotali_lead":
            st.subheader("Subtotali Lead")
            _render_bilancio_dataframe(
                _style_bilancio_table(
                    _scale_amount_columns(subtot_lead, amount_cols, amount_scale),
                    amount_cols,
                    amount_decimals
                )
            )
        elif selected_view == "gruppo_lead":
            st.subheader("Subtotali Gruppo Lead")
            _render_bilancio_dataframe(
                _style_bilancio_table(
                    _scale_amount_columns(subtot_group_lead, amount_cols, amount_scale),
                    amount_cols,
                    amount_decimals
                )
            )
        elif selected_view == "totali_tipo":
            st.subheader("Totali per tipo")
            _render_bilancio_dataframe(
                _style_bilancio_table(
                    _scale_amount_columns(subtot_tipo, amount_cols, amount_scale),
                    amount_cols,
                    amount_decimals
                )
            )

        # Esporta in Excel
        import io
        df_display_export = _prepare_export_dataframe(df_display, amount_cols, amount_scale, amount_decimals)
        df_display_no_account_export = _prepare_export_dataframe(df_display_no_account, amount_cols, amount_scale, amount_decimals)
        subtot_lead_export = _prepare_export_dataframe(subtot_lead, amount_cols, amount_scale, amount_decimals)
        subtot_group_lead_export = _prepare_export_dataframe(subtot_group_lead, amount_cols, amount_scale, amount_decimals)
        subtot_tipo_export = _prepare_export_dataframe(subtot_tipo, amount_cols, amount_scale, amount_decimals)
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df_display_export.to_excel(writer, index=False, sheet_name='Bilancio_Confronto')
            df_display_no_account_export.to_excel(writer, index=False, sheet_name='Bilancio_Senza_Conto')
            subtot_lead_export.to_excel(writer, index=False, sheet_name='Subtotali_Lead')
            subtot_group_lead_export.to_excel(writer, index=False, sheet_name='Subtotali_GroupLead')
            subtot_tipo_export.to_excel(writer, index=False, sheet_name='Subtotali_Tipo')
        excel_data = output.getvalue()
        col_export_excel, col_export_word, col_export_pdf = st.columns(3)
        with col_export_excel:
            st.download_button(
                label="Esporta in Excel",
                data=excel_data,
                file_name="bilancio_riepilogo.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        pdf_data = _build_pdf_by_lead(
            df_source=df_pivot,
            latest_col=latest_col,
            previous_col=previous_col,
            latest_year=latest_year,
            previous_year=previous_year,
            amount_scale=amount_scale,
            amount_decimals=amount_decimals,
        )
        if pdf_data is None:
            st.warning("Export PDF non disponibile: installare `reportlab`.")
        else:
            with col_export_pdf:
                st.download_button(
                    label="Esporta in PDF",
                    data=pdf_data,
                    file_name=f"bilancio_riepilogo_{latest_year}_vs_{previous_year}.pdf",
                    mime="application/pdf"
                )
        word_data = _build_docx_by_lead(
            df_source=df_pivot,
            latest_col=latest_col,
            previous_col=previous_col,
            latest_year=latest_year,
            previous_year=previous_year,
            amount_scale=amount_scale,
            amount_decimals=amount_decimals,
        )
        if word_data is None:
            st.warning("Export Word non disponibile: installare `python-docx`.")
        else:
            with col_export_word:
                st.download_button(
                    label="Esporta in Word",
                    data=word_data,
                    file_name=f"bilancio_riepilogo_{latest_year}_vs_{previous_year}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

except Exception as e:
    st.error("Errore nella pagina Bilancio Riepilogo.")
    st.exception(e)

finally:
    try:
        conn.close()
    except Exception:
        pass
